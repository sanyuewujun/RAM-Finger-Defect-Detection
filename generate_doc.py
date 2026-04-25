#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成系统介绍Word文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_system_doc():
    """创建系统介绍文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('金手指不良判断系统', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format = title.runs[0]
    title_format.font.size = Pt(24)
    title_format.font.bold = True
    
    # 副标题
    subtitle = doc.add_paragraph('基于深度学习的PCB金手指缺陷检测系统')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 空行
    
    # 系统概述
    doc.add_heading('1. 系统概述', level=1)
    overview = doc.add_paragraph(
        '本系统是一个基于深度学习神经网络的PCB金手指缺陷自动检测系统，'
        '旨在实现金手指不良品（NG）和良好品（OK）的高精度自动分类。'
        '系统采用先进的计算机视觉和深度学习技术，支持从模型训练、推理验证到结果管理的完整工作流程。'
    )
    
    # 核心功能
    doc.add_heading('2. 核心功能', level=1)
    
    # 2.1 数据处理
    doc.add_heading('2.1 数据处理与增强', level=2)
    data_feats = doc.add_paragraph(style='List Bullet')
    data_feats.add_run('自动数据加载与预处理：').bold = True
    data_feats.add_run('支持批量加载NG和OK分类的图像数据，自动进行数据集划分')
    
    data_feats2 = doc.add_paragraph(style='List Bullet')
    data_feats2.add_run('数据增强：').bold = True
    data_feats2.add_run('采用Albumentations库实现多种数据增强策略，提高模型泛化能力')
    
    data_feats3 = doc.add_paragraph(style='List Bullet')
    data_feats3.add_run('标准化处理：').bold = True
    data_feats3.add_run('对输入图像进行归一化、缩放等预处理操作，确保输入质量')
    
    # 2.2 模型训练
    doc.add_heading('2.2 模型训练', level=2)
    train_feats = doc.add_paragraph(style='List Bullet')
    train_feats.add_run('多模型支持：').bold = True
    train_feats.add_run('支持ResNet50和Swin Vision Transformer V2 B等多种预训练模型')
    
    train_feats2 = doc.add_paragraph(style='List Bullet')
    train_feats2.add_run('迁移学习：').bold = True
    train_feats2.add_run('利用ImageNet预训练权重，通过微调快速适配特定任务')
    
    train_feats3 = doc.add_paragraph(style='List Bullet')
    train_feats3.add_run('智能学习率调度：').bold = True
    train_feats3.add_run('采用OneCycleLR学习率策略，实现高效的模型收敛')
    
    train_feats4 = doc.add_paragraph(style='List Bullet')
    train_feats4.add_run('早停机制：').bold = True
    train_feats4.add_run('根据验证损失设置早停参数，防止过拟合')
    
    train_feats5 = doc.add_paragraph(style='List Bullet')
    train_feats5.add_run('模型检查点保存：').bold = True
    train_feats5.add_run('自动保存最佳模型权重和训练指标')
    
    # 2.3 推理与检测
    doc.add_heading('2.3 推理与检测', level=2)
    infer_feats = doc.add_paragraph(style='List Bullet')
    infer_feats.add_run('高效推理引擎：').bold = True
    infer_feats.add_run('支持GPU加速，快速进行单张或批量图像推理')
    
    infer_feats2 = doc.add_paragraph(style='List Bullet')
    infer_feats2.add_run('置信度输出：').bold = True
    infer_feats2.add_run('提供预测概率和置信度评分，便于质量控制')
    
    infer_feats3 = doc.add_paragraph(style='List Bullet')
    infer_feats3.add_run('实时推理：').bold = True
    infer_feats3.add_run('单张图像推理耗时在毫秒级，满足生产线实时检测需求')
    
    infer_feats4 = doc.add_paragraph(style='List Bullet')
    infer_feats4.add_run('自动分类与保存：').bold = True
    infer_feats4.add_run('推理结果自动分类保存至NG/OK文件夹，便于后续处理')
    
    # 2.4 结果管理
    doc.add_heading('2.4 结果管理与统计', level=2)
    result_feats = doc.add_paragraph(style='List Bullet')
    result_feats.add_run('历史记录管理：').bold = True
    result_feats.add_run('详细记录每次推理结果，包含时间戳、置信度、推理耗时等信息')
    
    result_feats2 = doc.add_paragraph(style='List Bullet')
    result_feats2.add_run('CSV导出功能：').bold = True
    result_feats2.add_run('支持将推理结果导出为CSV格式，便于数据分析和审计')
    
    result_feats3 = doc.add_paragraph(style='List Bullet')
    result_feats3.add_run('统计分析：').bold = True
    result_feats3.add_run('计算良率、缺陷率等关键指标，支持生成统计报告')
    
    # 2.5 用户界面
    doc.add_heading('2.5 图形化用户界面', level=2)
    ui_feats = doc.add_paragraph(style='List Bullet')
    ui_feats.add_run('PyQt5可视化界面：').bold = True
    ui_feats.add_run('提供友好的图形化操作界面，降低使用门槛')
    
    ui_feats2 = doc.add_paragraph(style='List Bullet')
    ui_feats2.add_run('多标签操作：').bold = True
    ui_feats2.add_run('支持推理、结果查看、模型管理等多个功能模块')
    
    ui_feats3 = doc.add_paragraph(style='List Bullet')
    ui_feats3.add_run('实时预览：').bold = True
    ui_feats3.add_run('支持在界面中实时预览检测图像和结果')
    
    ui_feats4 = doc.add_paragraph(style='List Bullet')
    ui_feats4.add_run('进度监控：').bold = True
    ui_feats4.add_run('显示推理进度条，实时反馈系统状态')
    
    ui_feats5 = doc.add_paragraph(style='List Bullet')
    ui_feats5.add_run('文件监控：').bold = True
    ui_feats5.add_run('支持实时监控指定文件夹，自动推理新增图像')
    
    # 2.6 命令行工具
    doc.add_heading('2.6 命令行工具', level=2)
    cli_feats = doc.add_paragraph(style='List Bullet')
    cli_feats.add_run('灵活调用：').bold = True
    cli_feats.add_run('支持命令行方式调用系统功能，便于集成到现有工作流程')
    
    cli_feats2 = doc.add_paragraph(style='List Bullet')
    cli_feats2.add_run('批量处理：').bold = True
    cli_feats2.add_run('支持批量处理图像文件夹，适合大规模检测任务')
    
    # 系统架构
    doc.add_heading('3. 系统架构', level=1)
    
    # 创建架构表
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '模块'
    header_cells[1].text = '功能描述'
    
    # 设置表头背景色
    for cell in header_cells:
        set_cell_background(cell, 'D3D3D3')
    
    # 填充表格
    rows_data = [
        ('数据处理模块', '数据加载、预处理、增强和dataset构建'),
        ('模型模块', '预训练模型选择、参数配置和网络架构'),
        ('训练模块', '模型训练、验证、早停、学习率调度'),
        ('推理模块', '加载模型权重、图像推理、置信度计算'),
        ('结果管理模块', '推理结果保存、统计分析、CSV导出'),
    ]
    
    for i, (module, desc) in enumerate(rows_data, 1):
        cells = table.rows[i].cells
        cells[0].text = module
        cells[1].text = desc
    
    doc.add_paragraph()  # 空行
    
    # 技术特点
    doc.add_heading('4. 技术特点', level=1)
    
    tech_feats = [
        ('深度学习框架', 'PyTorch框架，支持GPU加速计算'),
        ('预训练模型', '使用ImageNet预训练权重，快速收敛'),
        ('迁移学习', '采用微调策略，减少训练数据需求'),
        ('数据增强', 'Albumentations库实现多种增强策略'),
        ('性能优化', '支持混合精度训练、动态学习率调整'),
        ('易用性', 'GUI和CLI双界面支持，降低使用门槛'),
    ]
    
    for feature, detail in tech_feats:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature}: ').bold = True
        p.add_run(detail)
    
    doc.add_paragraph()  # 空行
    
    # 应用场景
    doc.add_heading('5. 应用场景', level=1)
    
    scenarios = doc.add_paragraph(style='List Bullet')
    scenarios.add_run('PCB金手指缺陷检测：')
    scenarios.add_run('在电子制造过程中自动检测金手指不良品，提高生产效率')
    
    scenarios2 = doc.add_paragraph(style='List Bullet')
    scenarios2.add_run('质量控制：')
    scenarios2.add_run('替代人工目视检查，降低成本，提高一致性')
    
    scenarios3 = doc.add_paragraph(style='List Bullet')
    scenarios3.add_run('工业应用：')
    scenarios3.add_run('可部署在生产线上进行实时检测和自动分类')
    
    doc.add_paragraph()  # 空行
    
    # 输出信息
    doc.add_heading('6. 系统输出', level=1)
    
    output_feats = [
        '预测类别（NG/OK）',
        '预测置信度（0-1概率值）',
        '单张图像推理耗时（毫秒）',
        '推理时间戳',
        '检测结果CSV记录',
        '分类后的图像文件',
    ]
    
    for item in output_feats:
        doc.add_paragraph(item, style='List Bullet')
    
    # 保存文档
    output_path = 'c:\\Users\\SatoY\\Desktop\\DeepLearning\\金手指不良判断系统功能介绍.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    create_system_doc()
